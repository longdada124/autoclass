import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import re
from datetime import datetime, timedelta

st.set_page_config(page_title="後龍國中課表暨調代課系統", layout="wide")

# --- 核心工具函數 ---
def master_replace(doc_obj, old_text, new_text):
    """替換 Word 內的文字，包含表格與段落"""
    new_val = str(new_text) if new_text is not None else ""
    # 替換段落
    for p in doc_obj.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    inline[i].text = inline[i].text.replace(old_text, new_val)
    # 替換表格
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old_text in p.text:
                        p.text = p.text.replace(old_text, new_val)

def get_week_dates(base_date):
    """計算該週週一至週五的日期字串"""
    start_of_week = base_date - timedelta(days=base_date.weekday())
    dates = []
    for i in range(5):
        d = start_of_week + timedelta(days=i)
        # 轉換為民國年格式 115.02.09
        roc_year = d.year - 1911
        dates.append(f"{roc_year}.{d.month:02d}.{d.day:02d}")
    return dates

def fill_sub_notice(template_bytes, teacher_name, changes, week_dates):
    """
    填寫代調課通知單
    changes: list of dict [{'day': 1-5, 'period': 1-8, 'text': '代702 國文'}]
    """
    doc = Document(BytesIO(template_bytes))
    
    # 1. 填寫標題姓名 (假設樣板中有 "周正軒" 或特定標記，這裡示範替換標題)
    # 建議您在 Word 樣板的名字旁加上 {{NAME}} 會更準確，目前先嘗試通用替換
    # 這裡假設您的樣板標題是 "苗栗縣立後龍國民中學代、調課通知單" 後面接名字
    # 我們嘗試直接在表格上方或標題找位置，或者替換掉範例名字 "周正軒"
    master_replace(doc, "周正軒", teacher_name)
    
    # 2. 填寫日期 (替換樣板上的範例日期)
    # 您的樣板有 115.02.09 ~ 115.02.13，我們依序替換
    # 為了避免誤判，建議您將樣板日期改為 {{D1}}, {{D2}}... 
    # 但若不改樣板，我們嘗試直接寫入表格第二列 (Row index 0 or 1 depending on header)
    
    # 鎖定第一個大表格
    if len(doc.tables) > 0:
        table = doc.tables[0]
        
        # 嘗試填寫日期：假設日期在 Header 的下一列或特定格
        # 根據您的截圖，日期在第一列(Row 0) 的 Column 3, 5, 7... 或 Row 1
        # 簡單作法：直接用 Week dates 填入對應格子 (需依照實際 Word 格子 index)
        # 這裡示範邏輯：
        # 週一日期: table.cell(0, 3).text = week_dates[0] (需視實際表格結構調整)
        pass 

    # 3. 填寫代調課內容
    # 根據您的截圖，表格左邊是 "1", "2"... 節次
    # 我們掃描表格第一欄來確認列數 (Row Index)
    row_map = {} # {'1': row_index, '2': row_index...}
    if len(doc.tables) > 0:
        table = doc.tables[0]
        for i, row in enumerate(table.rows):
            text = row.cells[0].text.strip()
            # 抓取全形或半形數字
            m = re.search(r'[1-9１-９]', text)
            if m:
                # 將全形轉半形以利對應
                num_map = {'１':'1','２':'2','３':'3','４':'4','５':'5','６':'6','７':'7','８':'8'}
                key = m.group()
                key = num_map.get(key, key)
                row_map[int(key)] = i

        # 開始填寫
        # 週一~週五對應的 Column Index (需根據您的 Word 表格實測)
        # 看截圖：週一=Col 2?, 週二=Col 4? (中間夾日期欄)
        # 假設結構：[星期, 時間, 週一, 日期, 週二, 日期...]
        # 索引推測：週一(2), 週二(4), 週三(6), 週四(8), 週五(10)
        day_col_map = {1: 2, 2: 4, 3: 6, 4: 8, 5: 10} 
        
        for chg in changes:
            r_idx = row_map.get(chg['period'])
            c_idx = day_col_map.get(chg['day'])
            if r_idx is not None and c_idx is not None:
                try:
                    # 避免 Index Error
                    if r_idx < len(table.rows) and c_idx < len(table.rows[r_idx].cells):
                        cell = table.cell(r_idx, c_idx)
                        # 保留原有換行，追加內容
                        cell.text = chg['text']
                except:
                    pass

    return doc

# --- 側邊欄 ---
with st.sidebar:
    st.header("⚙️ 系統資料管理")
    if st.button("🧹 重置所有資料"):
        st.session_state.clear()
        st.rerun()
    
    st.divider()
    st.info("請依序上傳三個檔案")
    f_assign = st.file_uploader("1. 配課表", type=["xlsx", "csv"])
    f_time = st.file_uploader("2. 課表", type=["xlsx", "csv"])
    f_sort = st.file_uploader("3. 教師排序表", type=["xlsx", "csv"])
    
    if f_assign and f_time and st.button("🚀 執行系統整合"):
        with st.spinner("正在讀取資料與樣板..."):
            try:
                # 1. 讀取 Excel
                df_assign = pd.read_excel(f_assign) if f_assign.name.endswith('xlsx') else pd.read_csv(f_assign)
                df_time = pd.read_excel(f_time) if f_time.name.endswith('xlsx') else pd.read_csv(f_time)
                
                # 2. 讀取 GitHub 內建樣板 (防呆)
                try:
                    with open("班級樣板.docx", "rb") as f: st.session_state.class_template = f.read()
                    with open("教師樣板.docx", "rb") as f: st.session_state.teacher_template = f.read()
                    with open("代調課通知單.docx", "rb") as f: st.session_state.sub_template = f.read()
                except FileNotFoundError:
                    st.warning("⚠️ 部分 Word 樣板未找到，請確認 GitHub 檔案名稱是否正確。")

                # 3. 解析資料核心邏輯
                assign_lookup = []
                all_teachers_db = set()
                tutors = {}
                
                # 解析配課
                for _, row in df_assign.iterrows():
                    c, s, t_raw = str(row['班級']).strip(), str(row['科目']).strip(), str(row['教師']).strip()
                    t_list = [name.strip() for name in t_raw.split('/')]
                    for t in t_list:
                        if t and t != "nan":
                            assign_lookup.append({'c': c, 's': s, 't': t})
                            all_teachers_db.add(t)
                    if s == "班級": tutors[c] = t_raw

                # 解析教師排序
                ordered_teachers = []
                base_hours = {}
                total_counts = {}
                all_teachers_list = list(all_teachers_db)
                
                if f_sort:
                    df_s = pd.read_excel(f_sort) if f_sort.name.endswith('xlsx') else pd.read_csv(f_sort)
                    for _, s_row in df_s.iterrows():
                        t_name = str(s_row.iloc[0]).strip()
                        if t_name in all_teachers_list:
                            ordered_teachers.append(t_name)
                            try: base_hours[t_name] = int(s_row.iloc[1])
                            except: base_hours[t_name] = 0
                    # 補上沒在排序表但在配課表的老師
                    for t in all_teachers_list:
                        if t not in ordered_teachers: ordered_teachers.append(t); base_hours[t] = 0
                else:
                    ordered_teachers = sorted(all_teachers_list)

                # 解析課表
                class_data = {}
                teacher_data = {}
                day_map = {"一":1,"二":2,"三":3,"四":4,"五":5,"週一":1,"週二":2,"週三":3,"週四":4,"週五":5}
                
                for _, row in df_time.iterrows():
                    c_raw, s_raw = str(row['班級']).strip(), str(row['科目']).strip()
                    d_str = str(row['星期']).strip()
                    d = day_map.get(d_str, 0)
                    p_match = re.search(r'\d+', str(row['節次']))
                    
                    if p_match and d > 0:
                        p = int(p_match.group())
                        # 找老師
                        curr_t_list = [x['t'] for x in assign_lookup if x['c'] == c_raw and x['s'] == s_raw]
                        display_t = "/".join(curr_t_list) if curr_t_list else "未知"
                        
                        # 填入班級資料
                        if c_raw not in class_data: class_data[c_raw] = {}
                        class_data[c_raw][(d, p)] = {"subj": s_raw, "teacher": display_t}
                        
                        # 填入教師資料
                        for t in curr_t_list:
                            if t not in teacher_data: teacher_data[t] = {}
                            teacher_data[t][(d, p)] = {"subj": s_raw, "class": c_raw}
                            total_counts[t] = total_counts.get(t, 0) + 1

                # 存入 Session
                st.session_state.update({
                    "class_data": class_data,
                    "teacher_data": teacher_data,
                    "tutors_map": tutors,
                    "base_hours": base_hours,
                    "total_counts": total_counts,
                    "ordered_teachers": ordered_teachers,
                    "data_ready": True
                })
                st.success("✅ 資料整合完畢！")
                st.rerun()

            except Exception as e:
                st.error(f"❌ 解析發生錯誤：{e}")

# --- 主畫面 ---
if st.session_state.get("data_ready"):
    tab1, tab2, tab3 = st.tabs(["🏫 班級課表", "👩‍🏫 教師課表", "📅 調代課管理(NEW)"])
    
    # --- 1. 班級課表 (還原功能) ---
    with tab1:
        classes = sorted(list(st.session_state.class_data.keys()))
        if not classes:
            st.warning("無班級資料")
        else:
            sel_c = st.selectbox("請選擇班級", classes)
            st.subheader(f"📍 {sel_c} 課表")
            
            # 建立表格數據
            c_rows = []
            for p in range(1, 9):
                row = {"節次": p}
                for d in range(1, 6):
                    info = st.session_state.class_data[sel_c].get((d,p))
                    txt = f"{info['subj']}\n{info['teacher']}" if info else ""
                    row[f"週{d}"] = txt
                c_rows.append(row)
            
            st.table(pd.DataFrame(c_rows).set_index("節次"))
            
            # 下載按鈕 (簡單版)
            if st.button("📥 下載此班級 Word") and st.session_state.get('class_template'):
                doc = Document(BytesIO(st.session_state.class_template))
                master_replace(doc, "{{CLASS}}", sel_c)
                # 簡單填入 (需配合樣板標籤 {{SD1P1}} 等)
                for d in range(1,6):
                    for p in range(1,9):
                        info = st.session_state.class_data[sel_c].get((d,p), {"subj":"","teacher":""})
                        master_replace(doc, f"{{{{SD{d}P{p}}}}}", info['subj'])
                        master_replace(doc, f"{{{{TD{d}P{p}}}}}", info['teacher'])
                buf = BytesIO()
                doc.save(buf)
                st.download_button("💾 下載檔案", buf.getvalue(), f"{sel_c}_課表.docx")

    # --- 2. 教師課表 (還原功能) ---
    with tab2:
        teachers = st.session_state.ordered_teachers
        sel_t = st.selectbox("請選擇教師", teachers)
        
        info_t = st.session_state.teacher_data.get(sel_t, {})
        base = st.session_state.base_hours.get(sel_t, 0)
        total = st.session_state.total_counts.get(sel_t, 0)
        
        st.write(f"**{sel_t}** 老師 | 應授: {base} | 實授: {total} | 兼代: {total - base}")
        
        t_rows = []
        for p in range(1, 9):
            row = {"節次": p}
            for d in range(1, 6):
                info = info_t.get((d,p))
                txt = f"{info['class']} {info['subj']}" if info else ""
                row[f"週{d}"] = txt
            t_rows.append(row)
        
        st.table(pd.DataFrame(t_rows).set_index("節次"))
        
        if st.button("📥 下載此教師 Word") and st.session_state.get('teacher_template'):
            doc = Document(BytesIO(st.session_state.teacher_template))
            master_replace(doc, "{{TEACHER}}", sel_t)
            # 填入
            for d in range(1,6):
                for p in range(1,9):
                    info = info_t.get((d,p), {"class":"","subj":""})
                    master_replace(doc, f"{{{{CD{d}P{p}}}}}", info['class'])
                    master_replace(doc, f"{{{{SD{d}P{p}}}}}", info['subj'])
            buf = BytesIO()
            doc.save(buf)
            st.download_button("💾 下載檔案", buf.getvalue(), f"{sel_t}_課表.docx")

    # --- 3. 調代課管理 (新功能) ---
    with tab3:
        st.header("🔄 調代課通知單產製")
        
        col1, col2 = st.columns(2)
        with col1:
            target_date = st.date_input("選擇代課日期", datetime.now())
            week_num = target_date.weekday() + 1
            week_dates = get_week_dates(target_date)
            st.caption(f"本週區間：{week_dates[0]} ~ {week_dates[4]}")

        with col2:
            absent_teacher = st.selectbox("請假/被代課教師", st.session_state.ordered_teachers, key="absent")
        
        # 顯示該師當日課程供選擇
        st.subheader("1. 選擇要代課的節次")
        day_lessons = []
        info_t = st.session_state.teacher_data.get(absent_teacher, {})
        
        for p in range(1, 9):
            info = info_t.get((week_num, p))
            if info:
                day_lessons.append({
                    "節次": p, 
                    "班級": info['class'], 
                    "科目": info['subj'],
                    "desc": f"第{p}節 - {info['class']}{info['subj']}"
                })
        
        if not day_lessons:
            st.warning(f"{absent_teacher} 老師在 {target_date} (週{week_num}) 沒有課程。")
        else:
            # 讓使用者選一節課
            selected_lesson = st.radio("請勾選課程：", day_lessons, format_func=lambda x: x['desc'])
            
            st.divider()
            st.subheader("2. 選擇代課教師 (自動推薦)")
            
            # 推薦邏輯：該時段空堂者
            available_teachers = []
            for t in st.session_state.ordered_teachers:
                # 檢查該老師在 (week_num, p) 是否有課
                if (week_num, selected_lesson['節次']) not in st.session_state.teacher_data.get(t, {}):
                    available_teachers.append(t)
            
            sub_teacher = st.selectbox("選擇代課老師 (已過濾空堂)", available_teachers)
            
            if st.button("🖨️ 產生代課通知單 (Word)"):
                if not st.session_state.get('sub_template'):
                    st.error("❌ 找不到【代調課通知單.docx】，請確認檔案。")
                else:
                    # 準備寫入資料
                    change_info = {
                        'day': week_num,
                        'period': selected_lesson['節次'],
                        'text': f"代{selected_lesson['班級']} {selected_lesson['科目']}"
                    }
                    
                    # 呼叫產製函數
                    doc_sub = fill_sub_notice(
                        st.session_state.sub_template,
                        sub_teacher, # 通知單是給代課老師的，所以抬頭寫代課老師
                        [change_info],
                        week_dates
                    )
                    
                    buf = BytesIO()
                    doc_sub.save(buf)
                    file_name = f"{target_date.strftime('%m%d')}_{sub_teacher}_代課單.docx"
                    st.download_button(f"⬇️ 下載 {sub_teacher} 的通知單", buf.getvalue(), file_name)
                    st.success(f"已生成給 {sub_teacher} 的通知單！")

else:
    st.info("👋 請於左側上傳 3 個資料檔並執行，系統將自動展開。")
